
#!/usr/bin/env python
# encoding: utf-8

'''
https://stackabuse.com/reading-and-writing-ms-word-files-in-python-via-python-docx-module/

pip install python-docx
'''
# from __future__ import (
#     absolute_import, division, print_function, unicode_literals
# )
import json
import copy
import locale
from turtle import pd
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from core import *

locale.setlocale(locale.LC_ALL, 'pl_PL.UTF-8')
locale._override_localeconv = {'mon_thousands_sep': ' '}

STYLE_ADDRESS = 'Adres'
STYLE_DATE = 'data'
STYLE_HEADER = 'wstep'
STYLE_GROUP = 'podsystem'
STYLE_ITEM = 'czesc'
STYLE_DESCR = 'opis'
STYLE_PARAM = 'parametr'
STYLE_SUMMARY = 'podsumowanie'
ID_PREFIX = '(art.'
ID_SUFFIX = ')'
RANGE_LEN = 3
CURRENCY = 'EUR'
CURRENCY_COST = 'EUR'

CLIENT_NAME = 'nazwa klienta'
CLIENT_ADDRESS = 'adres klienta'
CLIENT_ZIP_CODE = 'kod pocztowy klienta'
CLIENT_ADDRESSEE = 'imię nazwisko klienta'
INSTALLATION_PLACE = 'miejsce instalacji'
DATE = 'data'
OFFER_VALIDITY = 'termin ważności oferty'
DISCOUNT = 'rabat'

TOTAL = 'total'
ARGUMENTS = [
    CLIENT_NAME,
    CLIENT_ADDRESS,
    CLIENT_ZIP_CODE,
    CLIENT_ADDRESSEE,
    INSTALLATION_PLACE,
    DATE,
    OFFER_VALIDITY,
    DISCOUNT,
]

GROUPS_TOKEN = 'groups'
SUMMARY_TOKEN = 'summary'
FOOTER_TOKEN = 'footer'
PAYMENT_TOKEN = 'payment'
NONE_TOKEN = 'none'
INDEX_INCR = 2
OPEN_NEW_OFFER = 'Wybierz polską ofertę do przepisania:'
NAME_NEW_OFFER = 'Nazwij polską ofertę'
OPEN_PATTERN_OFFER = 'Wybierz polską ofertę do uaktualnienia bazy danych:'
DOCUMENT_MS_WORD = [('plik MS Word', 'doc'), ('plik MS Word', 'docx')]
GRUPA = 'Grupa'

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.

    https://stackoverflow.com/questions/43637211/retrieve-document-content-with-document-structure-with-python-docx
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

class WriteGroup:
    group_count = 0
    article_count = 2
    group = None
    groups = []

    def add_summary(table):
        total = 0
        for group in WriteGroup.groups:
            total += group.cost
            cells = table.rows[0].cells
            clear_table_row(cells)

            par = cells[0].add_paragraph()
            par.text = str(group.range[0])
            par.style = STYLE_SUMMARY
            par.alignment = 0 # 0 for left, 1 for center, 2 right, 3 justify
            
            par = cells[1].add_paragraph()
            par.text = '-'
            par.style = STYLE_SUMMARY
            par.alignment = 1 # 0 for left, 1 for center, 2 right, 3 justify            

            par = cells[2].add_paragraph()
            par.text = str(group.range[1] - INDEX_INCR)
            par.style = STYLE_SUMMARY
            par.alignment = 2

            par = cells[3].add_paragraph()
            par.text = group.title
            par.style = STYLE_SUMMARY
            par.alignment = 0

            par = cells[4].add_paragraph()
            par.text = CURRENCY_COST
            par.style = STYLE_ITEM
            par.alignment = 2

            par = cells[5].add_paragraph()
            par.text = locale.format_string(
                '%.2f', group.cost, grouping=True, monetary=True)
            par.style = STYLE_ITEM
            par.alignment = 2                          
    
            tbl = table._tbl
            tbl = copy.deepcopy(tbl)
            par = group.doc.add_paragraph()
            par._p.addnext(tbl)
        
        return total

    def __init__(self, number, title, doc):
        self.number = number
        WriteGroup.groups.append(self)
        # import pdb; pdb.set_trace()
        self.title = title
        self.range = [WriteGroup.article_count, WriteGroup.article_count]
        self.cost = 0
        self.doc = doc
        par = doc.add_paragraph()
        par.style = STYLE_GROUP
        par.text = f'{GRUPA} {self.number}: {title}'
        WriteGroup.group = self

    def add_article(self, table, article):
        cells = table.rows[0].cells
        clear_table_row(cells)
        self.cost += 0.0 if (article.is_optional or type(article.cost) != float) else article.cost
        par = cells[0].add_paragraph()
        par.text = '0' * (RANGE_LEN - len(str(self.range[1]))) + str(self.range[1]) + '.'
        par.style = STYLE_ITEM
        par.alignment = 2 # for left, 1 for center, 2 right, 3 justify
        WriteGroup.article_count += 2
        self.range[1] = WriteGroup.article_count
        
        par = cells[1].add_paragraph()
        par.text = f'{article.article_db.pl}\n{ID_PREFIX}{article.article_db.article_id}{ID_SUFFIX}'
        par.style = STYLE_ITEM
        par.alignment = 0 # 0 for left, 1 for center, 2 right, 3 justify

        par = cells[2].add_paragraph()
        par.text = f'{article.count} szt.'
        par.style = STYLE_ITEM
        par.alignment = 2

        par = cells[3].add_paragraph()
        par.text = CURRENCY
        par.style = STYLE_ITEM
        par.alignment = 2

        par = cells[4].add_paragraph()
        par.text = locale.format_string('%.2f', article.price, grouping=True, monetary=True) \
            if type(article.price) == float else article.price
        par.style = STYLE_ITEM
        par.alignment = 2

        par = cells[5].add_paragraph()
        par.text = CURRENCY_COST
        par.style = STYLE_ITEM
        par.alignment = 2 

        par = cells[6].add_paragraph()
        if article.is_optional:
            par.text = locale.format_string('%.2f', 0.0, grouping=True, monetary=True) + '\nopcja'
        else:
            par.text = locale.format_string('%.2f', article.cost, grouping=True, monetary=True) \
                if type(article.price) == float  else article.cost

        par.style = STYLE_ITEM
        par.alignment = 2 

        tbl = table._tbl
        tbl = copy.deepcopy(tbl)
        par = self.doc.add_paragraph()
        par._p.addnext(tbl)

    def add_article_params(self, table, params):
        # `table` is single-row one. Add rows: 
        for i in range(len(params) - 1):
            table.add_row()
        
        row = 0
        for _, param in params.items():
            cells = table.rows[row].cells
            row += 1
            clear_table_row(cells)

            par = cells[0].add_paragraph()
            par.text = param.pl + ':'
            par.style = STYLE_PARAM
            par.alignment = 0 # 0 for left, 1 for center, 2 right, 3 justify


            par = cells[1].add_paragraph()
            par.text = param.value
            par.style = STYLE_PARAM
            par.alignment = 0

        tbl = table._tbl
        tbl1 = copy.deepcopy(tbl)
        # import pdb; pdb.set_trace()        
        par = self.doc.paragraphs[len(self.doc.paragraphs) - 1]
        par._p.addnext(tbl1)

        # Cleaning, remove added rows:
        for i in range(1, len(table.rows)):
            tbl.remove(table.rows[1]._tr)

    def add_descr(self, descr):
        par = self.doc.add_paragraph()
        par.text = descr
        par.style = STYLE_DESCR
        par.alignment = 0 # 0 for left, 1 for center, 2 right, 3 justify


class OfferPl:  
    def __init__(self, args_json=None):       
        self.item_table = None
        self.parameter_table = None
        self.summary_table = None
        self.parameter_table = None
        self.args_json = args_json
        WriteGroup.article_count = 2
        
        if self.args_json is None:
            with open(CONFIG_JSON, 'r', encoding='utf-8') as f:
                self.args_json = json.load(f)
            
            for param in ARGUMENTS:
                if param == 'total':
                    continue

                value = input(
    f'''\n{param}: {self.args_json[param]} ?
    ENTER OK, tekst + ENTER - nowa wartość parametru:
    ''')
                if not value:
                    value = self.args_json[param]
                self.args_json[param] = value

        
            with open(CONFIG_JSON, 'w', encoding='utf-8') as f:
                json.dump(self.args_json, f, indent=4) 

        def new_docx(offer_file):
            doc = Document(TEMPLATE_DOC)
            # doc._body.clear_content()
            doc.save(offer_file)
            print('oferta: ' + offer_file)
            args_json[OFFER_PL] = offer_file
            with open(CONFIG_JSON, 'w', encoding='utf-8') as f:
                json.dump(args_json, f, indent=4)            
            return doc
        
        self.doc = open_offer_file(
            OFFER_PL, OPEN_NEW_OFFER, DOCUMENT_MS_WORD, new_docx, 
            args_json=self.args_json, new=True)
        if self.doc is None:
            return
        
        self.offer_pl = self.args_json[OFFER_PL]

        template = Document(TEMPLATE_DOC)
        self.item_table = template.tables[0]
        self.parameter_table = template.tables[1]        
        self.summary_table = template.tables[2]
        self.payment_table = template.tables[3]

        header = []
        self.groups = []
        self.summary = []
        self.payment = []
        self.footer = []

        paragraph_list = header
        _paragraph_list = paragraph_list

        for p in template.paragraphs:
            inline = p.runs
            text = ''
            for i in range(len(inline)):
                text += inline[i].text
            
            for param in ARGUMENTS:
                key = '${' + param + '}'
                if key in text:
                    text = text.replace(key, self.args_json[param])
                    inline[0].text = text
                    for i in range(1, len(inline)):
                        inline[i].text = ''

            if '${' + GROUPS_TOKEN + '}' in text:
                for i in range(len(inline)):
                    inline[i].text = ''
                _paragraph_list = self.groups
            elif '${' + SUMMARY_TOKEN + '}' in text:
                for i in range(len(inline)):
                    inline[i].text = ''
                _paragraph_list = self.summary
            elif '${' + PAYMENT_TOKEN + '}' in text:
                for i in range(len(inline)):
                    inline[i].text = ''
                _paragraph_list = self.payment                
            elif '${' + FOOTER_TOKEN + '}' in text:
                for i in range(len(inline)):
                    inline[i].text = ''
                _paragraph_list = self.footer
            elif '${' + NONE_TOKEN + '}' in text:
                for i in range(len(inline)):
                    inline[i].text = ''
                _paragraph_list = None                

            if paragraph_list == _paragraph_list:
                if paragraph_list is not None:
                    paragraph_list.append(p)
            paragraph_list = _paragraph_list

        self.doc._body.clear_content()
        # AttributeError: 'NoneType' object has no attribute '_body'

        for p in header:
            self.doc.add_paragraph(p.text, p.style)
        
    def add_group(self, index, title):
        par = self.doc.add_paragraph()
        par.text = ''        
        WriteGroup(index, title, self.doc)

    def add_comment(self, comment):
        pass     

    def add_article(self, article):
        WriteGroup.group.add_article(self.item_table, article)

    def add_article_params(self, params):
        WriteGroup.group.add_article_params(self.parameter_table, params)

    def add_descr(self, descr):
        WriteGroup.group.add_descr(descr)

    def finish(self):
        for _par in self.summary:
            par = self.doc.add_paragraph()
            par.style = STYLE_HEADER
            text = _par.text
            par.text = text

        price = WriteGroup.add_summary(self.summary_table)
        discount = float(self.args_json[DISCOUNT])
        total = locale.format_string(
            '%.2f', price - discount, grouping=True, monetary=True)

        # import pdb; pdb.set_trace()
        self.payment_table.rows[0].cells[2].paragraphs[0].text \
            = locale.format_string('%.2f', price, grouping=True, monetary=True)
        self.payment_table.rows[1].cells[2].paragraphs[0].text \
            = locale.format_string('%.2f', discount, grouping=True, monetary=True)
        self.payment_table.rows[2].cells[2].paragraphs[0].text = total

        tbl = self.payment_table._tbl
        par = self.doc.add_paragraph()
        par._p.addnext(tbl)

        key = '${' + TOTAL + '}'
        for _par in self.payment:
            text = _par.text
            if key in text:
                text = text.replace(key, total + ' ' + CURRENCY_COST)

            par = self.doc.add_paragraph()
            par.text = text

        for _par in self.footer:
            par = self.doc.add_paragraph()
            par.text = _par.text                        
        
    def save_offer(self):
        # for p in self.footer:
        #     self.doc.add_paragraph(p.text, p.style)
        try:
            self.doc.save(self.offer_pl)
        except Exception as ex:
            print(ex)
        self.doc = None


def main():
    OfferPl.read()
    # offer_pl = OfferPl()
    # offer_pl.add_group('Grupa pierwsza')
    # offer_pl.add_article(
    #     '21032708', 'Koszty montażu', 1, 65500.00, 65500.00)
    # offer_pl.add_article(
    #     '21033308', 'Rura', 1, 65500.00, 65500.00)
    # offer_pl.finish()
    # offer_pl.save_offer()


if __name__ == '__main__':
    main()

