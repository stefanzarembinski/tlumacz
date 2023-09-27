import sys
import re
import json
from docx import Document
from docx.text.paragraph import Paragraph
from core import *
from offer_pl import *
from offer_de import *

ARTICLE_RX = r'(?P<name>.+?) *\(*art\. *(?P<article_id>\d+) *\)*'

class Teacher:
    def __init__(self, args_json=None):
        self.offer_de = OfferDe(TEACH_DE, args_json)
        # import pdb; pdb.set_trace()
        self.offer_de.parse()

        # Reads a (translated) file ...
        self.doc = open_offer_file(
            TEACH_PL, OPEN_PATTERN_OFFER, DOCUMENT_MS_WORD, Document, args_json=args_json)

        self.groups_pl = []
        self.just_translated = {}

    def teach(self):
        """ Reads a translated file in order to enhance the dictionary, possibly.
        """
        if not self.offer_de.groups:
            return
        
        group = None
        style = None
        article = None
        article_id = None
        descr = None
        
        for block in iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                if not block.text:
                    continue
                _style = block.style.name
                if not style == _style:
                    debug(f'<{_style}>')
                debug(f'{block.text}')
                if _style == STYLE_GROUP:
                    group = Group()
                    self.groups_pl.append(group)
                    if not GRUPA in block.text:
                        raise HealthyException(
f"""
{block.text.strip()}

Powyższy paragraf tytułu grupy jest - chyba - źle sformatowany.

Tytuły grup muszą być jednym paragrafem - muszą być w jednej linii - (bez znaku ENTER) w stylu "podsystem". Gdy chcesz oddać wrażenie nowej linii, użyj SHIFT+ENTER, choć lepiej stosować automatyczne łamanie lini a zapobiegać nieładnym przypadkom przez użycie niełamiących się spacji CTRL+SHIFT+SPACE.
""")
                    elif not re.search(rf'{GRUPA} (\d+):', block.text):
                        raise HealthyException(
f"""
{block.text.strip()}

Powyższy tytułu grupy jest - chyba - źle napisany.

Tytuły grup muszą być zaczynać się słowem "{GRUPA}", potem numer, potem dwukropek.
""")                        

                    group.index = re.search(rf'{GRUPA} (\d+):', block.text)
                    group.index = int(group.index[1])
                    group.name = re.sub(rf'{GRUPA} \d+: ', '', block.text).strip()
                elif _style == STYLE_DESCR:
                    descr += block.text + '\n'
                    if article:
                        article.set_descr(pl=descr)

            else: # table
                # import pdb; pdb.set_trace()
                # Podstawa filtra MS 12/ 5 1550 mm h (art.394507)
                if not len(block.rows):
                    continue
                
                _style = block.rows[0].cells[0].paragraphs[0].style.name
                if not style == _style:
                    debug(f'<{_style}>')
                
                if _style == STYLE_ITEM:
                    for row in block.rows:
                        _article = ''
                        for par in row.cells[1].paragraphs:
                            _article += ' ' + par.text
                        _article = _article.strip()
                        debug(_article)
                        
                        if ID_PREFIX in _article:
                            # import pdb; pdb.set_trace()
                            _article = re.search(ARTICLE_RX, _article)                          
                            name = _article['name']
                            article_id = _article['article_id']
                            # if article_id == '000048980':
                            #     import pdb; pdb.set_trace()                        

                            article = None
                            descr = ''
                            con = sqlite3.connect(DB)
                            cur = con.cursor()
                                                          

                            if ID_SUFFIX in article_id:
                                article_id = article_id[:article_id.index(ID_SUFFIX)].strip()
                            debug(f'{article_id}: {name}')
                            article = Article(article_id, de=None, pl=name)
                            group.articles.append(article)

                            if article.article_db.is_regex:
                                # article is not represented in the database (regex article?)!
                                article = None
                            else:
                                _de = article.article_db.de
                                if article.article_db.is_translated == JUST_TRANSLATED:
                                    self.just_translated[(article_id, solidHash(_de))] \
                                        = (_de, article.article_db.pl)
                                article = article.set_descr(pl='')
                                descr = ''

                elif _style == STYLE_PARAM:
                    for row in block.rows:
                        name = ''
                        for par in row.cells[0].paragraphs:
                            name += ' ' + par.text
                        name = re.sub(rf'\s*:\s*$', '', name)
                        value = ''
                        for par in row.cells[1].paragraphs: 
                            value += ' ' + par.text
                        debug(name)
                        if article:
                            article.set_param(value=value , pl=name)
                
            style = _style
        
        for group_pl in self.groups_pl:        
            db_group(self.offer_de.angebot_nr, group_pl.index, group_pl.name, translate=False)             

        articles_json = {}
        con = sqlite3.connect(DB)
        cur = con.cursor()
        cur.execute(f"SELECT * FROM {ARTICLES}")
        _articles = cur.fetchall()
        for _article in _articles:
            _article_id = _article[0]
            _article_hash = _article[1]
            # if _article_id == '375420':
            #     import pdb; pdb.set_trace()
            _article = pickle.loads(_article[2])
            articles_json[str((_article_id, _article_hash))] = json.loads(
                json.dumps(_article.__dict__, default=lambda o: o.__dict__))

        with open(ARTICLES_JSON, 'w', encoding='utf-8') as f:
            json.dump(articles_json, f, indent=4, sort_keys=True)

        if self.just_translated:
            indexes = sorted(self.just_translated)
            print(f'\n### Elementy dodane do tłumacza:')
            for id in indexes:
                names = self.just_translated[id]
                print(f'{id[0]}: "{names[0]}" => "{names[1]}"')
            print()
    
def main():
    teacher = Teacher()
    teacher.teach()

if __name__ == '__main__':
    main()